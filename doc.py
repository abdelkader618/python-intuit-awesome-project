# to be able to run this program type pip3 install -r requirements.txt in terminal
from docx import Document 
from docx.shared import Cm
import pyttsx3 # pyttsx3 enables us to turn text to voice
def speak(text):
    pyttsx3.speak (text)
speak('salam alikom ya mizyeana ')
document=Document()

# profile picture
document.add_picture(
    'me.jpg',
    width=Cm(12.0)# use widht and height to control the size of the picture
    )# use .add_  to add stuff to your document 
# name
name =  input ('whats your name ')
speak('hello  '+ name + 'how are you fine i hope')
speak('give me your phone number punk')
phone_number =input ('whats your phone number '  )
email =input ('whats your email ')
document.add_paragraph (
    name + ' | ' + phone_number + ' | ' + email )
# about me
document.add_heading('about me ')
about_me = input ('tell me about yourself ')
speak('stop flattering yourself ')
document.add_paragraph(about_me)
#  work experience
document.add_heading('work experience ')
p=document.add_paragraph()
company = input('enter company ')
from_date=input('from date ')
to_date=input('to date ')
p.add_run(company +' ').bold=True #use add_run to add text to an existing paragraph control the font by using .bold and .italic
p.add_run(from_date+' - '+to_date + '\n').italic=True #use '\n'to add a new line to ypur paragraph
experience_details=input(
    'describe your experience at ' + company
)
p.add_run(experience_details)
#more experiences
while True :
    has_more_experiences = input ('do you have more experiences yes or no ')
    if has_more_experiences.lower()=='yes':#use .lower() for small caracters and .upper() for capital caracters
        p=document.add_paragraph()
        company = input('enter company ')
        from_date=input('from date ')
        to_date=input('to date ')
        p.add_run(company +' ').bold=True #use add_run to add text to an existing paragraph control the font by using .bold and .italic
        p.add_run(from_date+' - '+to_date + '\n').italic=True #use '\n'to add a new line to ypur paragraph
        experience_details=input(
            'describe your experience at ' + company +' '
        )
        p.add_run(experience_details )
    else:
        break
# skills 
document.add_heading('skills that i have ')
skils=input('enter skills ')
p=document.add_paragraph (skils)
p.style ='List Bullet' # make sure you wrote the style type correctly List Bullet not list bullet
while True : 
    has_more_skills= input ('do u have more skills yes or no ')
    if has_more_skills.lower()=='yes' :
        skils=input('enter skills ')
        p= document.add_paragraph (skils)
        p.style = 'List Bullet'
    else: 
        break
# footer
section =document.sections[0]#to add a footer you need access to a section 
footer=section.footer # then put the section into an object
p=footer.paragraphs[0]# then invoke the paragraoh function
p.text='CV generated using some really awesome skills'
document.save('cv.docx')